#!/usr/bin/env python3
"""Build original synthetic DOCX references for style and audit forward-testing."""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SKILL_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_SPEC = SKILL_ROOT / "assets" / "samples" / "coursepack-spec.synthetic.json"
DEFAULT_TOKENS = SKILL_ROOT / "assets" / "templates" / "docx-style-tokens.json"


def load_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_run(run, tokens: dict[str, Any], *, size: float, bold=False, italic=False, color=None):
    fonts = tokens["fonts"]
    colors = tokens["colors"]
    latin = fonts["latin"]
    east = fonts["cjk_fallback"]
    run.font.name = latin
    r_fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)
    r_fonts.set(qn("w:eastAsia"), east)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = rgb(color or colors["ink"])
    return run


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top=55, bottom=55, start=90, end=90):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.find(qn("w:tcMar"))
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for edge, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def cell_border(cell, **edges):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge_name, attrs in edges.items():
        node = borders.find(qn(f"w:{edge_name}"))
        if node is None:
            node = OxmlElement(f"w:{edge_name}")
            borders.append(node)
        for key, value in attrs.items():
            node.set(qn(f"w:{key}"), str(value))


def table_widths(table, widths: list[int]):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[min(index, len(widths) - 1)]
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell_margins(cell)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def paragraph_bottom_border(paragraph, color: str):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "5")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def configure(doc: Document, tokens: dict[str, Any], role: str):
    page = tokens["page"]
    section = doc.sections[0]
    section.page_width = Inches(page["width_inches"])
    section.page_height = Inches(page["height_inches"])
    section.left_margin = Inches(page["margin_left_inches"])
    section.right_margin = Inches(page["margin_right_inches"])
    section.top_margin = Inches(page["margin_top_inches"])
    section.bottom_margin = Inches(page["margin_bottom_inches"])
    section.header_distance = Inches(page["header_distance_inches"])
    section.footer_distance = Inches(page["footer_distance_inches"])

    normal = doc.styles["Normal"]
    normal.font.name = tokens["fonts"]["latin"]
    normal._element.rPr.rFonts.set(qn("w:ascii"), tokens["fonts"]["latin"])
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), tokens["fonts"]["latin"])
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), tokens["fonts"]["cjk_fallback"])
    normal.font.size = Pt(tokens["fonts"]["student_body_pt"] if role == "student" else tokens["fonts"]["teacher_body_pt"])
    normal.paragraph_format.space_after = Pt(2)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(header.add_run("Reading Explorer · Synthetic Reference"), tokens, size=8.2, bold=True, color=tokens["colors"]["gray"])
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)
    set_run(run, tokens, size=8.2, color=tokens["colors"]["gray"])


def page_break(doc: Document):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.page_break_before = True
    paragraph.paragraph_format.space_after = Pt(0)


def add_student_title(doc: Document, spec: dict[str, Any], tokens: dict[str, Any]):
    colors = tokens["colors"]
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run("Reading Explorer Foundations"), tokens, size=9.5, bold=True, color=colors["green"])
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    set_run(p.add_run(spec["title"]), tokens, size=19, bold=True, color=colors["green_dark"])
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    set_run(p.add_run("Practice Book"), tokens, size=12.5, bold=True, color=colors["gold"])
    set_run(p.add_run("  |  Vocabulary · Sentences · Text Building · Review"), tokens, size=10, color=colors["gray"])
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    set_run(p.add_run("Name 姓名  ____________________    Class 班级  ______________    Date 日期  ______________"), tokens, size=9.5, bold=True, color=colors["gray"])


def add_page_title(doc: Document, title: str, cn: str, tokens: dict[str, Any]):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(4)
    set_run(p.add_run(f"{title}  |  {cn}"), tokens, size=14.5, bold=True, color=tokens["colors"]["green_dark"])


def add_exercise_heading(doc: Document, number: int, title: str, cn: str, tokens: dict[str, Any]):
    table = doc.add_table(rows=1, cols=1)
    table_widths(table, [10420])
    cell = table.cell(0, 0)
    shade_cell(cell, tokens["colors"]["green"])
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run(f"Exercise {number} · {title}  |  {cn}"), tokens, size=10.7, bold=True, color=tokens["colors"]["white"])


def add_instruction(doc: Document, exercise: dict[str, Any], tokens: dict[str, Any]):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(3)
    set_run(p.add_run(exercise["instruction_en"]), tokens, size=9.7, bold=True, color=tokens["colors"]["gold"])
    cn = exercise.get("instruction_cn")
    if cn:
        set_run(p.add_run(f"  {cn}"), tokens, size=9.2, color=tokens["colors"]["gray"])


def add_word_bank(doc: Document, words: str, tokens: dict[str, Any], fill=None):
    table = doc.add_table(rows=1, cols=1)
    table_widths(table, [10420])
    shade_cell(table.cell(0, 0), fill or tokens["colors"]["gold_pale"])
    cell_border(table.cell(0, 0), top={"val": "single", "sz": 6, "color": tokens["colors"]["gold"]}, bottom={"val": "single", "sz": 6, "color": tokens["colors"]["gold"]}, start={"val": "single", "sz": 6, "color": tokens["colors"]["gold"]}, end={"val": "single", "sz": 6, "color": tokens["colors"]["gold"]})
    p = table.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run(words), tokens, size=9.5, bold=True, color=tokens["colors"]["green_dark"])


def add_items(doc: Document, lines: list[str], tokens: dict[str, Any], *, answer_line=False, size=9.8):
    for index, line in enumerate(lines, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.first_line_indent = Inches(-0.10)
        p.paragraph_format.space_after = Pt(1 if not answer_line else 0)
        set_run(p.add_run(f"{index}.  {line}"), tokens, size=size)
        if answer_line:
            lp = doc.add_paragraph()
            lp.paragraph_format.left_indent = Inches(0.32)
            lp.paragraph_format.space_after = Pt(2)
            paragraph_bottom_border(lp, tokens["colors"]["line"])
            set_run(lp.add_run("\u00A0"), tokens, size=8.5, color=tokens["colors"]["gray"])


def exercise_lookup(spec: dict[str, Any]) -> dict[str, dict[str, Any]]:
    result: dict[str, dict[str, Any]] = {}
    for section in spec["practice_book"]["sections"]:
        for exercise in section["exercises"]:
            result[exercise["exercise_id"]] = exercise
    return result


CN_TITLES = {
    "PB01": "英英释义",
    "PB02": "固定搭配",
    "PB03": "语境选词",
    "PB04": "选择正确词形",
    "PB05": "发现并修正错误",
    "PB06": "合并信息",
    "PB07": "句型转换",
    "PB08": "重组文章发展顺序",
    "PB09": "概要完形填空",
    "PB10": "综合选择",
    "PB11": "判断证据",
    "PB12": "四句报告",
}


def add_exercise(doc: Document, number: int, exercise: dict[str, Any], tokens: dict[str, Any]):
    exercise_id = exercise["exercise_id"]
    add_exercise_heading(doc, number, exercise["title"], CN_TITLES[exercise_id], tokens)
    add_instruction(doc, exercise, tokens)

    if exercise_id == "PB01":
        rows = [
            ("signal", "a message or sign"),
            ("direction", "the way that something moves or points"),
            ("theory", "an idea that tries to explain facts"),
            ("possibility", "something that may be true or may happen"),
            ("evidence", "information that supports a claim"),
            ("explain", "to make something clear"),
        ]
        definitions = [rows[2][1], rows[4][1], rows[1][1], rows[5][1], rows[0][1], rows[3][1]]
        table = doc.add_table(rows=7, cols=4)
        table_widths(table, [550, 1900, 900, 7070])
        headers = ["No.", "Word", "Answer", "English Definitions"]
        for index, text in enumerate(headers):
            shade_cell(table.cell(0, index), tokens["colors"]["green"])
            p = table.cell(0, index).paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run(p.add_run(text), tokens, size=9.1, bold=True, color=tokens["colors"]["white"])
        repeat_header(table.rows[0])
        for row_index, (word, _) in enumerate(rows, 1):
            cells = table.rows[row_index].cells
            values = [str(row_index), word, "____", f"{chr(96 + row_index)}. {definitions[row_index - 1]}"]
            if row_index % 2 == 0:
                for cell in cells:
                    shade_cell(cell, tokens["colors"]["gray_pale"])
            for col, value in enumerate(values):
                p = cells[col].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col in {0, 2} else WD_ALIGN_PARAGRAPH.LEFT
                set_run(p.add_run(value), tokens, size=8.9)
    elif exercise_id == "PB02":
        add_word_bank(doc, "another · change · explain · send · supporting · theory", tokens)
        add_items(doc, [
            "____________________ direction",
            "____________________ a signal",
            "suggest a ____________________",
            "____________________ possibility",
            "____________________ evidence",
            "____________________ the change",
        ], tokens, size=9.4)
    elif exercise_id == "PB03":
        add_word_bank(doc, "direction · exactly · perhaps · signal · surprising · sure", tokens, tokens["colors"]["blue_pale"])
        add_items(doc, [
            "The sensor recorded something ____________________ near the gate.",
            "The model car changed ____________________ after the turn.",
            "The flashing light may be a ____________________.",
            "Nobody is ____________________ ____________________ why the alarm stopped.",
            "____________________ the battery failed, but the evidence is incomplete.",
        ], tokens, size=9.3)
    elif exercise_id == "PB04":
        add_items(doc, [
            "The result was ____________________.  (surprise)",
            "The students were ____________________ by the result.  (surprise)",
            "The light moved ____________________.  (strange)",
            "Nobody knows ____________________ what caused it.  (exact)",
            "This could ____________________ the change.  (explanation)",
            "All possibilities should be ____________________.  (consider)",
        ], tokens)
    elif exercise_id == "PB05":
        add_items(doc, [
            "The camera saw a strangely light near the gate.",
            "The light's direction was change after the turn.",
            "This could explains the second signal.",
            "Another possibilities should be discussed.",
            "All evidence should be consider carefully.",
        ], tokens, answer_line=True, size=9.4)
    elif exercise_id == "PB06":
        p = doc.add_paragraph()
        set_run(p.add_run("Example  "), tokens, size=9.2, bold=True, color=tokens["colors"]["green_dark"])
        set_run(p.add_run(exercise["example"]), tokens, size=9.1)
        add_items(doc, [
            "The camera turned on. A light was moving across the field.  (when)",
            "They had seen reflections before. This light was different.  (before, but)",
            "The light passed the window. Its direction changed.  (after + -ing)",
            "Its direction changed twice. The class suggested another theory.  (therefore)",
            "The drone cannot be seen. The camera information remains useful.  (but)",
        ], tokens, answer_line=True, size=9.3)
    elif exercise_id == "PB07":
        add_items(doc, [
            "The class named the light 'Green Line.' → The light ________________________________.",
            "We should consider every possibility. → Every possibility ________________________.",
            "The camera can no longer see it. → It ____________________________________________.",
            "simple theory / model drone → The simplest theory is that ________________________.",
            "perhaps / reflection / window → __________________________________________________.",
        ], tokens, answer_line=True, size=9.2)
    elif exercise_id == "PB08":
        add_word_bank(doc, "alternative theory · first explanation · new evidence · observation · result · unknown point", tokens, tokens["colors"]["blue_pale"])
        statements = [
            "A  Some learners therefore suggested a different theory.",
            "B  A camera recorded a long light crossing the field.",
            "C  Nobody is sure exactly what produced the light.",
            "D  Later, the light changed direction twice.",
            "E  The simplest idea was that a model drone caused it.",
            "F  Another learner said a window reflection was possible.",
        ]
        for line in statements:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.12)
            p.paragraph_format.space_after = Pt(0)
            set_run(p.add_run(line), tokens, size=9.0)
        table = doc.add_table(rows=7, cols=3)
        table_widths(table, [850, 1900, 7670])
        for col, text in enumerate(("Step", "Sentence Letter", "Function")):
            shade_cell(table.cell(0, col), tokens["colors"]["green"])
            p = table.cell(0, col).paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run(p.add_run(text), tokens, size=9, bold=True, color=tokens["colors"]["white"])
        repeat_header(table.rows[0])
        for row in range(1, 7):
            if row % 2 == 0:
                for cell in table.rows[row].cells:
                    shade_cell(cell, tokens["colors"]["gray_pale"])
            for col, value in enumerate((str(row), "____", "____________________________")):
                p = table.cell(row, col).paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col < 2 else WD_ALIGN_PARAGRAPH.LEFT
                set_run(p.add_run(value), tokens, size=8.8)
    elif exercise_id == "PB09":
        add_word_bank(doc, "camera · direction · evidence · mystery · possibility · signal · sure · theory · window", tokens)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.08)
        p.paragraph_format.line_spacing = 1.22
        summary = (
            "A school (1) __________________ recorded a long light. Its speed and (2) __________________ changed. "
            "Some learners thought it was a (3) __________________ from a drone. The simplest (4) __________________ "
            "was a model aircraft, but the new (5) __________________ did not prove it. A reflection was another "
            "(6) __________________. Nobody is (7) __________________ exactly what happened, so the event remains a "
            "(8) __________________."
        )
        set_run(p.add_run(summary), tokens, size=9.3)
    elif exercise_id == "PB10":
        add_items(doc, [
            "The camera saw something ___.  a surprise   b surprising   c surprisingly",
            "A long light ___ across the field.  a moved   b was moving   c had moved",
            "They ___ reflections before that night.  a saw   b were seeing   c had seen",
            "The camera recorded a ___ signal.  a strange   b strangely   c strangeness",
            "The light was ___ shaped.  a strange   b strangely   c strangeness",
            "Its direction changed after ___ the wall.  a pass   b passed   c passing",
            "This could ___ the result.  a explain   b explains   c explained",
            "The drone can no longer ___ from the field.  a see   b be seen   c seen",
        ], tokens, size=8.8)
    elif exercise_id == "PB11":
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.08)
        p.paragraph_format.right_indent = Inches(0.08)
        set_run(p.add_run("The Blue Flash  "), tokens, size=9.1, bold=True, color=tokens["colors"]["green_dark"])
        set_run(p.add_run("At 8:10, a gate camera recorded a blue flash. It moved upward and disappeared. The simplest idea is that a bicycle light caused it. Perhaps a phone screen reflected in the glass. Nobody is sure exactly what the camera recorded."), tokens, size=8.9)
        add_items(doc, [
            "_____  The camera recorded a blue flash.",
            "_____  The flash stayed in one place.",
            "_____  The bicycle-light explanation has been proved.",
            "_____  A phone reflection is another possibility.",
            "_____  A witness saw the bicycle directly.",
        ], tokens, size=8.8)
        p = doc.add_paragraph()
        set_run(p.add_run("Repair the three unsupported claims:"), tokens, size=9.0, bold=True, color=tokens["colors"]["green_dark"])
        table = doc.add_table(rows=3, cols=2)
        table_widths(table, [380, 10040])
        for row_index, row in enumerate(table.rows, 1):
            set_run(row.cells[0].paragraphs[0].add_run(f"{row_index}."), tokens, size=8.7, color=tokens["colors"]["gray"])
            cell_border(row.cells[1], bottom={"val": "single", "sz": 5, "color": tokens["colors"]["line"]})
            set_run(row.cells[1].paragraphs[0].add_run("\u00A0"), tokens, size=8.7)
    elif exercise_id == "PB12":
        labels = [
            "Evidence: The camera recorded ...",
            "Simplest idea: The simplest idea is that ...",
            "Another possibility: Perhaps ...",
            "Unknown point: Nobody is sure exactly ...",
        ]
        for label in labels:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            set_run(p.add_run(label), tokens, size=9.0, bold=True, color=tokens["colors"]["green_dark"])
            lp = doc.add_paragraph()
            lp.paragraph_format.left_indent = Inches(0.15)
            lp.paragraph_format.space_after = Pt(2)
            paragraph_bottom_border(lp, tokens["colors"]["line"])
            set_run(lp.add_run("\u00A0"), tokens, size=8.5)


def build_practice_book(spec: dict[str, Any], tokens: dict[str, Any], path: Path):
    doc = Document()
    configure(doc, tokens, "student")
    add_student_title(doc, spec, tokens)
    exercises = exercise_lookup(spec)
    pages = [
        ("Vocabulary and Collocations", "词汇与固定搭配", ["PB01", "PB02", "PB03"]),
        ("Words in Use", "词形与准确运用", ["PB04", "PB05"]),
        ("Sentences and Grammar", "句子与语法", ["PB06", "PB07"]),
        ("Text Building", "语篇结构与概要", ["PB08", "PB09"]),
        ("Unit Review", "单元综合练习", ["PB10", "PB11", "PB12"]),
    ]
    for page_index, (title, cn, refs) in enumerate(pages):
        if page_index > 0:
            page_break(doc)
        add_page_title(doc, title, cn, tokens)
        for ref in refs:
            add_exercise(doc, int(ref[2:]), exercises[ref], tokens)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def add_teacher_title(doc: Document, spec: dict[str, Any], tokens: dict[str, Any]):
    colors = tokens["colors"]
    p = doc.add_paragraph()
    set_run(p.add_run("Reading Explorer Foundations"), tokens, size=10, bold=True, color=colors["green"])
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    set_run(p.add_run(spec["title"]), tokens, size=21, bold=True, color=colors["green_dark"])
    p = doc.add_paragraph()
    set_run(p.add_run("Teacher Unit Guide"), tokens, size=13, bold=True, color=colors["gold"])
    set_run(p.add_run("  |  Complete Lesson-by-Lesson Design"), tokens, size=10, color=colors["gray"])
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    set_run(p.add_run("This synthetic reference demonstrates the approved structure only. Replace all content from verified teacher-owned sources in a private project."), tokens, size=9.2, color=colors["gray"])
    add_page_title(doc, "Unit Route", "单元路线", tokens)
    route = spec["unit_route"]
    table = doc.add_table(rows=6, cols=2)
    table_widths(table, [3800, 6620])
    rows = [
        ("Reading A · Deep Reading", f"{route['reading_a_deep_reading']} lessons"),
        ("Reading B · Transfer Reading", f"{route['reading_b_transfer_reading']} lessons"),
        ("Video / Listening", f"{route['video_listening']} lessons"),
        ("Integrated Output", f"{route['integrated_output']} lesson"),
        ("Retrieval / Feedback", f"{route['retrieval_feedback']} lesson"),
        ("Default artifacts", "Teacher Unit Guide + Student Practice Book"),
    ]
    for row_index, values in enumerate(rows):
        if row_index % 2 == 0:
            for cell in table.rows[row_index].cells:
                shade_cell(cell, tokens["colors"]["green_pale"])
        for col, value in enumerate(values):
            set_run(table.cell(row_index, col).paragraphs[0].add_run(value), tokens, size=9.0, bold=col == 0)


def add_teacher_label(doc: Document, label: str, body: str, tokens: dict[str, Any]):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    set_run(p.add_run(f"{label}: "), tokens, size=9.1, bold=True, color=tokens["colors"]["green_dark"])
    set_run(p.add_run(body), tokens, size=9.0)


def add_lesson(doc: Document, number: int, lesson: dict[str, Any], tokens: dict[str, Any]):
    page_break(doc)
    colors = tokens["colors"]
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    set_run(p.add_run(f"Lesson {number} · {lesson['title']}"), tokens, size=16.5, bold=True, color=colors["green_dark"])
    add_teacher_label(doc, "Purpose 目的", lesson["purpose"], tokens)
    add_teacher_label(doc, "Success evidence 成功证据", lesson["success_evidence"], tokens)
    add_teacher_label(doc, "Source route 来源与练习", " | ".join(lesson["source_route"]), tokens)
    add_teacher_label(doc, "Materials 材料", lesson["materials"], tokens)
    add_teacher_label(doc, "Preparation 课前准备", lesson["preparation"], tokens)

    add_page_title(doc, "Lesson Flow", "40分钟完整教学流程", tokens)
    for stage_index, stage in enumerate(lesson["stages"], 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(1)
        set_run(p.add_run(f"{stage_index}. {stage['title']}  ({stage['minutes']} min · {stage['grouping']})"), tokens, size=10.2, bold=True, color=colors["green_dark"])
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.12)
        p.paragraph_format.space_after = Pt(1)
        set_run(p.add_run(stage["teacher_move_cn"]), tokens, size=9.0)
        if stage.get("ready_to_say_en"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(1)
            set_run(p.add_run("Ready to say: "), tokens, size=8.9, bold=True, color=colors["gold"])
            set_run(p.add_run(stage["ready_to_say_en"]), tokens, size=8.9, italic=True)
        add_teacher_label(doc, "Student action 学生行动", stage["student_action"], tokens)
        add_teacher_label(doc, "Product / evidence 产出与个体证据", f"{stage['product']} | {stage['individual_evidence']}", tokens)
        add_teacher_label(doc, "Check / response 检查与预期", f"{stage['check']} | {stage['anticipated_response']}", tokens)
        add_teacher_label(doc, "Likely error / recovery 易错与补救", f"{stage['likely_error']} | {stage['recovery_move']}", tokens)
        add_teacher_label(doc, "Board / transition 板书与过渡", f"{stage['board_display']} | {stage['transition']}", tokens)

    add_page_title(doc, "After the Flow", "课后决策", tokens)
    add_teacher_label(doc, "Board plan 板书方案", lesson["board_plan"], tokens)
    add_teacher_label(doc, "Anticipated errors 常见问题", " | ".join(lesson["anticipated_errors"]), tokens)
    add_teacher_label(doc, "Access support 支持", lesson["differentiation"]["access"], tokens)
    add_teacher_label(doc, "Extension 提升", lesson["differentiation"]["extension"], tokens)
    if lesson.get("activities"):
        for activity in lesson["activities"]:
            add_teacher_label(doc, "Optional activity 可选活动", activity["title"], tokens)
            add_teacher_label(doc, "Rules 规则", activity["rules"], tokens)
            add_teacher_label(doc, "Scoring / evidence 评分与证据", f"{activity['teacher_adjudication']} | {activity['individual_evidence']}", tokens)
            add_teacher_label(doc, "Fallback 低准备替代", activity["fallback"], tokens)
    add_teacher_label(doc, "Lesson close 结课", lesson["lesson_close"], tokens)
    homework = lesson["homework"]
    add_teacher_label(doc, "Homework 作业", f"{homework['title']} | Practice Book {', '.join(homework['exercise_refs'])} | {homework['minutes']} min", tokens)
    add_teacher_label(doc, "Teacher verification 教师核查", homework["teacher_verification"], tokens)
    add_teacher_label(doc, "Next-lesson use 次课使用", homework["next_lesson_use"], tokens)


def build_teacher_guide(spec: dict[str, Any], tokens: dict[str, Any], path: Path):
    doc = Document()
    configure(doc, tokens, "teacher")
    add_teacher_title(doc, spec, tokens)
    for index, lesson in enumerate(spec["lessons"], 1):
        add_lesson(doc, index, lesson, tokens)
    page_break(doc)
    add_page_title(doc, "Answer Key and Judgment", "答案与判断标准", tokens)
    table = doc.add_table(rows=len(spec["teacher_answers"]) + 1, cols=2)
    table_widths(table, [1500, 8920])
    for col, text in enumerate(("Answer Ref", "Answer or Acceptable Response")):
        shade_cell(table.cell(0, col), tokens["colors"]["green"])
        set_run(table.cell(0, col).paragraphs[0].add_run(text), tokens, size=9.0, bold=True, color=tokens["colors"]["white"])
    repeat_header(table.rows[0])
    for row_index, (answer_ref, answer) in enumerate(spec["teacher_answers"].items(), 1):
        if row_index % 2 == 0:
            for cell in table.rows[row_index].cells:
                shade_cell(cell, tokens["colors"]["gray_pale"])
        set_run(table.cell(row_index, 0).paragraphs[0].add_run(answer_ref), tokens, size=8.4, bold=True)
        set_run(table.cell(row_index, 1).paragraphs[0].add_run(str(answer)), tokens, size=8.4)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--spec", type=Path, default=DEFAULT_SPEC)
    parser.add_argument("--tokens", type=Path, default=DEFAULT_TOKENS)
    parser.add_argument("--out", type=Path, required=True)
    args = parser.parse_args()
    try:
        spec = load_json(args.spec)
        tokens = load_json(args.tokens)
        practice = args.out / "Synthetic_Student_Practice_Book.docx"
        teacher = args.out / "Synthetic_Teacher_Unit_Guide.docx"
        build_practice_book(spec, tokens, practice)
        build_teacher_guide(spec, tokens, teacher)
    except (OSError, KeyError, ValueError, json.JSONDecodeError) as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        return 1
    print(practice)
    print(teacher)
    return 0


if __name__ == "__main__":
    sys.exit(main())
